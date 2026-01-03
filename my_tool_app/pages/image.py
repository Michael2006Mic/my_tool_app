import streamlit as st
st.set_page_config(page_title="Image to Word (AI)", page_icon="🖼️")
import easyocr
import numpy as np
from PIL import Image
from docx import Document
import io

# ==========================================
# CONFIGURATION
# ==========================================
st.set_page_config(page_title="Image to Word (AI)", page_icon="🖼️")

# ==========================================
# LOGIC: EASYOCR (No Tesseract needed)
# ==========================================

@st.cache_resource
def load_reader():
    """
    Load the AI model once and keep it in memory.
    Using CPU is safer for free cloud servers to avoid crashing.
    """
    return easyocr.Reader(['en'], gpu=False) 

def image_to_word_easyocr(image_file):
    output_filename = "extracted_text.docx"
    
    try:
        # 1. Load Image
        image = Image.open(image_file)
        
        # Convert to numpy array (EasyOCR expects this)
        image_np = np.array(image)

        # 2. Extract Text
        reader = load_reader()
        
        # detail=0 gives us just the text strings
        result = reader.readtext(image_np, detail=0, paragraph=True)
        
        if not result:
            return None

        # 3. Create Word Doc
        doc = Document()
        doc.add_heading('Extracted Text (EasyOCR)', 0)
        
        # EasyOCR returns a list of paragraph strings
        for paragraph in result:
            doc.add_paragraph(paragraph)
            
        doc.save(output_filename)
        return output_filename
        
    except Exception as e:
        st.error(f"Error: {e}")
        return None

# ==========================================
# UI
# ==========================================

st.title("🖼️ AI Image to Word")
st.markdown("Extract text from images . No installation required.")

uploaded_file = st.file_uploader("Upload Image", type=["png", "jpg", "jpeg"])

if uploaded_file is not None:
    st.image(uploaded_file, caption="Uploaded Image", width=300)
    
    if st.button("Convert to Word"):
        with st.spinner("Initializing AI and scanning... (First time takes 10s)"):
            result_file = image_to_word_easyocr(uploaded_file)
            
            if result_file:
                st.success("Extraction Complete!")
                with open(result_file, "rb") as f:
                    st.download_button(
                        label="Download Word Doc",
                        data=f,
                        file_name="scanned_text.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("No text found in this image.")